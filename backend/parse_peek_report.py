"""Utilities for parsing PEEK energy reports into structured organ/chakra data."""

from __future__ import annotations

import re
from pathlib import Path
from typing import IO, Any, Protocol, cast


class DocumentLike(Protocol):
    paragraphs: list[Any]
    tables: list[Any]


class DocumentLoader(Protocol):
    def __call__(self, docx: str | IO[bytes] | None = None) -> DocumentLike: ...


Document: DocumentLoader | None
_document_import_error: Exception | None

try:
    from docx import Document as _RealDocument
except Exception as exc:  # pragma: no cover - runtime dependency
    Document = None
    _document_import_error = exc
else:
    Document = cast(DocumentLoader, _RealDocument)
    _document_import_error = None

# ---------------------------------------------------------------------------
# Constants & mappings
# ---------------------------------------------------------------------------

ARROW_SEPARATORS = r"(->|→|➔|➜)"
CHEVR_SEPARATORS = "([>\\u203A])"
VALUE_RE = re.compile(r"(\d{1,3})\b")
VALUE_AFTER_CHEVRON = re.compile("(?<!-)[>\\u203A]\\s*(\\d{1,3})\\b")
PAREN_LABEL_RE = re.compile(r"\(([^)]+)\)")

TARGET_ORGAN_IDS = {
    "brain",
    "thyroid",
    "lungs",
    "heart",
    "lymphatic",
    "liver",
    "spleen",
    "stomach",
    "gallbladder",
    "kidneys",
    "small_intestine",
    "large_intestine",
    "bladder",
    "reproductive_male",
    "reproductive_female",
}

NAME_TO_ID = {
    "brain": "brain",
    "br": "brain",
    "thyroid": "thyroid",
    "th": "thyroid",
    "lung": "lungs",
    "lungs": "lungs",
    "lu": "lungs",
    "heart": "heart",
    "hr": "heart",
    "pericardium": "heart",
    "pc": "heart",
    "lymphatic": "lymphatic",
    "ly": "lymphatic",
    "liver": "liver",
    "lv": "liver",
    "spleen": "spleen",
    "sp": "spleen",
    "stomach": "stomach",
    "st": "stomach",
    "gallbladder": "gallbladder",
    "gall bladder": "gallbladder",
    "gb": "gallbladder",
    "kidney": "kidneys",
    "kidneys": "kidneys",
    "ki": "kidneys",
    "small intestine": "small_intestine",
    "si": "small_intestine",
    "large intestine": "large_intestine",
    "li": "large_intestine",
    "bladder": "bladder",
    "bl": "bladder",
    "pancreas": "pancreas_placeholder",
    "pa": "pancreas_placeholder",
    "reproductive": "__reproductive__",
    "re": "__reproductive__",
}

CHAKRA_ID_BY_NUM = {
    1: "Chakra_01_Root",
    2: "Chakra_02_Sacral",
    3: "Chakra_03_SolarPlexus",
    4: "Chakra_04_Heart",
    5: "Chakra_05_Throat",
    6: "Chakra_06_ThirdEye",
    7: "Chakra_07_Crown",
}

SUPPORTED_CHAKRA_IDS = set(CHAKRA_ID_BY_NUM.values())

SPECIAL_METRIC_KEYS = {
    "inflammatory score": "inflammatory_score",
    "immunal defense": "immunal_defense",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _tokenize_line(raw: str) -> list[str]:
    if not raw:
        return []
    normalized = raw.replace("\u200b", " ").strip()
    if not normalized:
        return []
    normalized = re.sub(r"[\t•]+", " ", normalized)
    normalized = re.sub(ARROW_SEPARATORS, "|", normalized)
    normalized = re.sub(CHEVR_SEPARATORS, "|", normalized)
    parts = [part.strip(" -") for part in normalized.split("|") if part.strip(" -")]
    if len(parts) <= 1:
        parts = [token.strip() for token in re.split(r"\s{2,}", normalized) if token.strip()]
    if len(parts) <= 1:
        parts = [token.strip() for token in normalized.split() if token.strip()]
    return parts


def _extract_value(tokens: list[str]) -> tuple[int | None, int | None]:
    for idx in range(len(tokens) - 1, -1, -1):
        match = VALUE_RE.search(tokens[idx])
        if match:
            return int(match.group(1)), idx
    return None, None


def _norm(text: str) -> str:
    text = text.strip().lower()
    text = re.sub(r"[.\u200b]", " ", text)
    return re.sub(r"\s+", " ", text)


def _map_name_to_id(name: str) -> str | None:
    key = _norm(name)
    if key in NAME_TO_ID:
        return NAME_TO_ID[key]
    base = re.sub(r"\s*\(.*?\)\s*", "", key).strip()
    if base in NAME_TO_ID:
        return NAME_TO_ID[base]
    for segment in re.split(r"[-/]", base):
        segment = segment.strip()
        if segment in NAME_TO_ID:
            return NAME_TO_ID[segment]
    return None


def _parse_organ_tokens(tokens: list[str]) -> tuple[str, int] | None:
    if not tokens or not tokens[0].lower().startswith("organ"):
        return None

    body = tokens[1:]
    if not body:
        return None

    value, value_idx = _extract_value(body)
    if value is None:
        return None

    name_tokens = body[:value_idx]
    if not name_tokens:
        return None

    # remove leading short uppercase codes (e.g. LI, SI)
    filtered = list(name_tokens)
    while filtered and re.fullmatch(r"[A-Za-z]{1,4}[.:]?", filtered[0]):
        filtered.pop(0)
    if not filtered:
        filtered = name_tokens

    candidates: list[str] = []
    for start in range(len(filtered)):
        candidate = " ".join(filtered[start:]).strip()
        if candidate:
            candidates.append(candidate)
    candidates.extend(name_tokens)

    for candidate in candidates:
        organ_id = _map_name_to_id(candidate)
        if organ_id:
            return organ_id, value

    return None


def _parse_chakra_tokens(tokens: list[str]) -> tuple[str, int] | None:
    if not tokens or not tokens[0].lower().startswith("chakra"):
        return None

    body = tokens[1:]
    if not body:
        return None

    value, value_idx = _extract_value(body)
    if value is None or value_idx is None:
        return None

    number: int | None = None
    for idx in range(value_idx + 1):
        match = re.search(r"(\d+)", body[idx])
        if match:
            number = int(match.group(1))
            break

    if number is None or number not in CHAKRA_ID_BY_NUM:
        return None

    return CHAKRA_ID_BY_NUM[number], value


def _value_after_chevron(raw_line: str) -> int | None:
    match = VALUE_AFTER_CHEVRON.search(raw_line)
    if not match:
        return None
    return int(match.group(1))


def _parse_organ_strict(raw_line: str) -> tuple[str, int] | None:
    if not raw_line.lower().startswith("organs"):
        return None
    value = _value_after_chevron(raw_line)
    if value is None:
        return None
    tokens = _tokenize_line(raw_line)
    parsed = _parse_organ_tokens(tokens)
    if not parsed:
        return None
    organ_id, _ = parsed
    return organ_id, value


def _parse_chakra_strict(raw_line: str) -> tuple[str, int] | None:
    if not raw_line.lower().startswith("chakra"):
        return None
    value = _value_after_chevron(raw_line)
    if value is None:
        return None
    tokens = _tokenize_line(raw_line)
    parsed = _parse_chakra_tokens(tokens)
    if not parsed:
        return None
    chakra_id, _ = parsed
    return chakra_id, value


def _parse_special_metric(
    tokens: list[str],
    raw_line: str,
) -> tuple[str, dict[str, Any]] | None:
    value, value_idx = _extract_value(tokens)
    if value is None or value_idx is None or value_idx <= 1:
        return None

    name_tokens = tokens[1:value_idx]
    filtered = [
        token
        for token in name_tokens
        if not re.fullmatch(r"[A-Za-z]{1,3}(?:[.:])?", token) or _norm(token) in SPECIAL_METRIC_KEYS
    ]
    if filtered:
        name_tokens = filtered

    if not name_tokens:
        return None

    metric_name = _norm(" ".join(name_tokens))
    metric_id = SPECIAL_METRIC_KEYS.get(metric_name)
    if not metric_id:
        return None

    label_match = PAREN_LABEL_RE.search(raw_line)
    label = label_match.group(1).strip() if label_match else None
    display_name = " ".join(name_tokens).strip()

    metric: dict[str, Any] = {"value": value}
    if display_name:
        metric["name"] = display_name
    if label:
        metric["label"] = label

    return metric_id, metric


def _read_doc_all_lines(path: Path) -> list[str]:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from _document_import_error

    doc = Document(str(path))
    lines: list[str] = []

    for paragraph in doc.paragraphs:
        value = (paragraph.text or "").strip()
        if value:
            lines.append(value)

    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            first = row.cells[0]
            text = " ".join(
                (p.text or "").strip() for p in first.paragraphs if (p.text or "").strip()
            )
            text = text.strip()
            if text:
                lines.append(text)

    return lines


def read_word_lines(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix != ".docx":
        raise RuntimeError(
            f"Unsupported file type for PEEK report: {path.suffix}. Expected '.docx'."
        )
    return _read_doc_all_lines(path)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_report(path: Path) -> dict[str, Any]:
    organs: dict[str, int] = {}
    chakras: dict[str, int] = {}
    metrics: dict[str, dict[str, Any]] = {}

    for raw_line in read_word_lines(path):
        tokens = _tokenize_line(raw_line)
        if not tokens:
            continue

        header = tokens[0].lower()
        if header.startswith("organ"):
            parsed_strict = _parse_organ_strict(raw_line)
            if parsed_strict:
                organ_id, value = parsed_strict
                if organ_id == "__reproductive__":
                    organs.setdefault("reproductive_male", value)
                    organs.setdefault("reproductive_female", value)
                elif organ_id == "pancreas_placeholder":
                    continue
                elif organ_id in TARGET_ORGAN_IDS:
                    organs[organ_id] = value
                continue

            metric_parsed = _parse_special_metric(tokens, raw_line)
            if metric_parsed:
                metric_id, metric_values = metric_parsed
                metrics[metric_id] = metric_values
            continue

        if header.startswith("chakra"):
            parsed_strict = _parse_chakra_strict(raw_line)
            if parsed_strict:
                chakra_id, value = parsed_strict
                if chakra_id in SUPPORTED_CHAKRA_IDS:
                    chakras[chakra_id] = value
            continue

    result: dict[str, Any] = {"organs": organs, "chakras": chakras}
    if metrics:
        result["metrics"] = metrics
    return result


def main() -> None:  # pragma: no cover - CLI helper
    import argparse
    import json

    parser = argparse.ArgumentParser(description="Parse a PEEK Word report into JSON.")
    parser.add_argument("input", help="Path to a .docx file")
    parser.add_argument("-o", "--out", default="cardenergymap_values.json", help="Output JSON path")
    args = parser.parse_args()

    data = parse_report(Path(args.input))
    with open(args.out, "w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False, indent=2)
    print(f"Wrote {args.out}")


if __name__ == "__main__":  # pragma: no cover
    main()
